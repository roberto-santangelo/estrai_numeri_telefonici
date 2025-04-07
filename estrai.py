# pyinstaller --onefile --name "Estrai numeri telefonici.exe" "estrai.py" --icon=logo.ico
import requests
from bs4 import BeautifulSoup
from collections import defaultdict
import re, os
import docx
from docx.shared import Inches, Cm, Pt, RGBColor
from getpass import getpass
import socket


def internet_connection():
    print("STO TESTANDO LA TUA CONNESSIONE AD INTERNET...")
    remote_server = "www.google.com"
    port = 80
    sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    sock.settimeout(5)
    try:
        sock.connect((remote_server, port))
        return True
    except socket.error:
        return False
    finally:
        print("\033[F" + "\033[2K\r", end="")
        sock.close()


if not internet_connection():
    print("NON SEI CONNESSO AD INTERNET. IL PROGRAMMA VIENE TERMINATO")
    x = getpass("PREMI INVIO PER CHIUDERE IL PROGRAMMA...")
    exit()

try:
    print("ESTRAI NUMERI DI TELEFONO v. 1.0 - by Roberto Santangelo\n ")
    indirizzi = []
    comune_input = input(
        "INSERISCI IL COMUNE/CITTA' DI CUI VUOI CERCARE GLI INDIRIZZI: "
    )
    while comune_input == "":
        comune_input = input("DEVI DIGITARE IL NOME DEL COMUNE: ")
    indirizzo = None
    i = 1
    msg = "INSERISCI UN INDIRIZZO ALLA VOLTA, QUANDO HAI FINITO LASCIA IL CAMPO DI INDIRIZZO VUOTO E PREMI INVIO"
    print(f"{'-'*(len(msg)+4)}\n| {msg} |\n{'-'*(len(msg)+4)}")
    while True:
        indirizzo = input(f"INSERISCI IL {i}° INDIRIZZO: ").strip()
        if (
            indirizzo == ""
        ):  # Se l'utente preme Invio senza inserire nulla, si esce dal ciclo
            if not indirizzi:
                print("DEVI INSERIRE ALMENO UN INDIRIZZO.")
                continue  # Torna all'inizio del ciclo
            print(
                "\033[F" + "\033[2K\r", end=""
            )  # Sposta il cursore su una riga più in alto + CANCELLO LA RIGA CORRENTE
            break  # Esce dal ciclo se c'è almeno un indirizzo
        indirizzi.append(indirizzo)
        i += 1  # Incrementa il contatore
    print("")
    elenco_numeri = set()  # Set per evitare duplicati più efficientemente
    numeri_di_telefono = defaultdict(list)  # Dizionario con i risultati

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }

    def pulisci_numero(numero):
        """Rimuove spazi e caratteri speciali, lasciando solo le cifre."""
        return re.sub(r"\D", "", numero)  # Rimuove tutto tranne i numeri

    n = 0
    total_results = 0  # Inizializza la variabile per evitare ReferenceError

    if indirizzi:
        for ind in indirizzi:
            indirizzo_encoded = requests.utils.quote(ind)

            # --- Pagine Bianche ---
            nome = via = telefono = località = None
            url = f"https://www.paginebianche.it/cerca-da-indirizzo?dv={comune_input.replace(' ', '%20')}%20{indirizzo_encoded}"
            # print(f"PAGINA VISITATA: {url}")
            print(
                f'STO CERCANDO SU PAGINE BIANCHE PER L\'INDIRIZZO "{ind} - {comune_input}" (PAGINA 1)'
            )
            response = requests.get(url, headers=headers)

            if response.status_code == 200:
                soup = BeautifulSoup(response.text, "html.parser")
                elementi = soup.find_all("h2", class_="list-element__title")

                for elem in elementi:
                    nome = elem.get_text(strip=True) or "Nome non disponibile"

                    # Estrazione indirizzo
                    indirizzo_tag = elem.find_next(
                        "div", class_="list-element__address"
                    )
                    if indirizzo_tag:
                        spans = indirizzo_tag.find_all("span")
                        via_completa = (
                            spans[0].get_text()
                            if len(spans) > 0
                            else "Indirizzo non trovato"
                        )
                        via_completa = via_completa.split(sep="-")
                        via = via_completa[0].strip()
                        località = via_completa[1].strip()
                    else:
                        via, località = "Indirizzo non trovato", "Località non trovata"

                    # Estrazione telefono
                    telefono_tag = elem.find_next("a", class_="phone-numbers__main")
                    telefono = (
                        pulisci_numero(telefono_tag.get_text(strip=True))
                        if telefono_tag
                        else "Telefono non trovato"
                    )
                    cap = località.split(" ", 1)[0]
                    provincia = località.rsplit(" ", 1)[-1]
                    comune = località.replace(cap, "").replace(provincia, "").strip()

                    match = re.search(r"\((.*?)\)", provincia)
                    provincia = match.group(1)
                    località = f"{cap} {comune} {provincia}"

                    if (
                        telefono not in elenco_numeri and telefono.isdigit()
                    ):  # Evita numeri errati
                        elenco_numeri.add(telefono)
                        numeri_di_telefono["PAGINE BIANCHE"].append(
                            {
                                "Nome": nome,
                                "Indirizzo": via,
                                "Telefono": telefono,
                                "Località": località,
                            }
                        )
            else:
                print(f"Errore {response.status_code} nell'accesso a {url}")

            # --- InElenco ---
            nome = via = telefono = località = None
            pagine = 1
            da = 0
            while (
                da <= (pagine * 10) or da > 1000
            ):  # CICLO PER CONTROLLARE TUTTE LE PAGINE
                url = f"https://mobile.inelenco.com/?dir=cerca&nome=&comune={comune_input.replace(' ', '+')}&provincia=&indirizzo={ind.replace(' ', '+')}&telefono=&da={da}"
                # print(f"PAGINA VISITATA: {url}")
                print(
                    f'STO CERCANDO SU INELENCO PER L\'INDIRIZZO "{ind} - {comune_input}" (PAGINA {(da//10)+1})'
                )
                response = requests.get(url, headers=headers)

                if response.status_code == 200:
                    soup = BeautifulSoup(response.text, "html.parser")

                    # Trova tutti i blocchi di risultati
                    entries = soup.find_all("td", id="risultati", class_="cerca")

                    if entries:
                        for entry in entries:
                            nome = entry.get_text(strip=True)  # Nome dell'attività
                            if re.match(r"^La ricerca ", nome):
                                continue
                            telefono = (
                                entry.find_next("td", class_="dativ")
                                .text.replace("Telefono", "")
                                .strip()
                            )
                            via = entry.find_next("td", class_="dati").text.strip()

                            # Trova il resto delle informazioni
                            dati = entry.find_all_next("td", class_="dati", limit=4)
                            comune = dati[1].text.replace("Comune", "").strip()
                            provincia = dati[2].text.replace("Provincia", "").strip()
                            cap = dati[3].text.replace("CAP", "").strip()

                            # match = re.search(r"\((.*?)\)", provincia)
                            # provincia = match.group(1) if match else provincia

                            if (
                                telefono not in elenco_numeri and telefono.isdigit()
                            ):  # Evita numeri errati
                                elenco_numeri.add(telefono)
                                numeri_di_telefono["INELENCO"].append(
                                    {
                                        "Nome": nome,
                                        "Indirizzo": via,
                                        "Telefono": telefono,
                                        "Località": f"{cap} {comune} {provincia}",
                                    }
                                )

                    if da == 0:  # CONTROLLO IL NUMERO DI PAGINE UNA VOLTA SOLA
                        # CONTROLLO SE CI SONO ALTRE PAGINE DI CUI ESTRARRE I NUMERI
                        tag_risultati = soup.find("td", id="risultati")
                        if tag_risultati:
                            match = re.search(r"de (\d+)", tag_risultati.get_text())
                            if match:
                                total_results = int(match.group(1))
                        pagine = total_results // 10

                else:
                    print(f"Errore {response.status_code} nell'accesso a {url}")

                da += 10  # VADO ALLA PAGINA SUCCESSIVA
                # time.sleep(0.1)  # Pausa di 1 secondo tra le richieste
    print("")
    # CERCO IL PREFISSO
    url = f"https://www.nonsolocap.it/cap?k={comune_input .replace(' ', '+')}"
    response = requests.get(url, headers=headers)

    if response.status_code == 200:
        soup = BeautifulSoup(response.text, "html.parser")
        prefisso = soup.find("td", class_="d-none d-md-table-cell")
        if prefisso:
            prefisso = prefisso.get_text(strip=True)
        else:
            prefisso = ""

    numero_contatti = (
        sum([len(v) for v in numeri_di_telefono.values()]) if numeri_di_telefono else 0
    )

    # --- Stampa risultati ---
    if numeri_di_telefono:
        doc = docx.Document()

        # CAMBIO LE SEZIONI
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)

        # CAMBIO GLI STILI DEI CARATTERI
        ## TITOLO
        stile_titolo = doc.styles["Heading 1"].font
        stile_titolo.name = "Calibri"
        stile_titolo.size = docx.shared.Pt(12)
        stile_titolo.color.rgb = RGBColor(53, 79, 82)

        ## STILE NORMALE
        stile_normale = doc.styles["Normal"]
        stile_normale.paragraph_format.line_spacing = 1
        font_normal = stile_normale.font
        font_normal.name = "Calibri"
        font_normal.size = Pt(12)

        formato_paragrafo = stile_normale.paragraph_format
        formato_paragrafo.space_before = Pt(0)
        formato_paragrafo.space_after = Pt(0)

        # Aggiungi paragrafo con spaziatura personalizzata
        # paragrafo = doc.add_paragraph("Testo di esempio con spaziatura personalizzata.")
        # paragrafo.paragraph_format.space_before = Pt(3)
        # paragrafo.paragraph_format.space_after = Pt(5)

        h1 = doc.add_heading(f"CONTATTI TERRITORIO - {cap} {comune.upper()}")

        for sito, contatti in numeri_di_telefono.items():
            paragrafo_sito = doc.add_paragraph(f"{sito}\n{'-' * 30}")
            for item in contatti:
                print(sito)
                print(f"Nome: {item['Nome']}")
                print(f"Indirizzo: {item['Indirizzo']}")
                print(f"Località: {item['Località']}")
                print(f"Telefono: {item['Telefono'].replace(prefisso, f'{prefisso} ')}")
                print("")

                paragraph = doc.add_paragraph(f"{item['Nome']}")
                paragraph = doc.add_paragraph(f"{item['Indirizzo']}")
                # document.add_paragraph(f"{item['Località']}")
                paragraph = doc.add_paragraph(
                    f"{item['Telefono'].replace(prefisso, f'{prefisso} ')}"
                )
                paragraph = doc.add_paragraph(f"")
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(5)

        print(f"NUMERO DI CONTATTI TROVATI: {numero_contatti}")

        nome_file = "contatti estratti.docx"
        salva = False
        if os.path.exists(nome_file):
            while True:
                risposta = (
                    input(
                        f'IL FILE "{nome_file}" È GIÀ ESISTENTE. VUOI SOVRASCRIVERLO? (S/N) '
                    )
                    .strip()
                    .upper()
                )

                if risposta == "S":
                    salva = True
                    break
                elif risposta == "N":
                    break
                else:
                    print("Risposta non valida. Inserisci 'S' per sì o 'N' per no.")
        else:
            salva = True

        if salva:
            doc.save(nome_file)
            print(f'Contatti salvati in "{nome_file}" nella cartella:\n{os.getcwd()}')

except Exception as err:
    print(f"ERRORE: {err}")
finally:
    getpass("PREMI INVIO PER CHIUDERE IL PROGRAMMA...")
